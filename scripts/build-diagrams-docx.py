"""
build-diagrams-docx.py — Generate the AeroAdmin AFM diagrams documentation
as a .docx using python-docx.

Input: 7 PNGs at docs/diagrams/assets/
Output: docs/diagrams/AeroAdmin-AFM-Diagramas.docx

Structure:
- Cover page
- Summary table
- 7 sections (one per diagram) with title, description, embedded image
- Appendix A: Reusable image paths
- Appendix B: How to extend the set (pointer to HANDOFF.md)
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

# -----------------------------------------------------------------------------
# Constants
# -----------------------------------------------------------------------------
REPO_ROOT = Path(__file__).resolve().parent.parent
ASSETS = REPO_ROOT / "docs" / "diagrams" / "assets"
OUT = REPO_ROOT / "docs" / "diagrams" / "AeroAdmin-AFM-Diagramas.docx"

# AeroAdmin brand skin (matches docs/diagrams/style-guide.md)
BRAND_GREEN = RGBColor(0x0B, 0x5F, 0x2D)  # accent
INK = RGBColor(0x1A, 0x2E, 0x22)  # ink dark forest
MUTED = RGBColor(0x4F, 0x5D, 0x75)  # muted blue-slate
SOFT = RGBColor(0x7A, 0x83, 0x99)  # soft
ALERT_RED = RGBColor(0xA9, 0x32, 0x32)  # concern

# Diagram catalog: (number, slug, title, type, focus, description, wired_in)
DIAGRAMS = [
    (
        1, "01-system-architecture", "System Architecture",
        "architecture",
        "Postgres+PostGIS (DB como corazón del sistema)",
        "Flujo end-to-end del sistema en 3 zonas: DJI Sources (drone + SmartFarm + kr-ag2-api), Scraper & Pipeline (Playwright client + JSON exports + run-pipeline.js), y AeroAdmin App (Postgres+PostGIS focal + V0 adapter + Next.js + User). Reemplaza el ASCII art de `docs/ARCHITECTURE.md §1`.",
        "docs/ARCHITECTURE.md §1 (callout al HTML)",
    ),
    (
        2, "02-dji-data-pipeline", "DJI Data Pipeline",
        "architecture (process-style)",
        "Paso 4 — Spatial join (la única transformación no-trivial)",
        "Los 9 pasos idempotentes de `scripts/run-pipeline.js` en 2 filas (5+4). El wrap-around dashed conecta el fin de la fila 1 con el inicio de la fila 2. Cada paso muestra el script que lo ejecuta y el artefacto que escribe (JSON file o tabla PostGIS).",
        "docs/ARCHITECTURE.md §1 + docs/DJI_SCRAPER.md TL;DR",
    ),
    (
        3, "03-fumigation-cadence-state", "Fumigation Cadence State Machine",
        "state",
        "Estado overdue (alerta, en concern red) + flecha de recovery (en accent green)",
        "Máquina de estados de cadencia: `no_history → ok → due_soon → overdue` con la flecha de recovery `overdue → ok` (accent verde). Cada estado muestra el guard (`diffDays = (now - next_due_date) / 86_400_000`) y los self-loops `tick` (tiempo que pasa sin fumigación nueva).",
        "docs/FUMIGATION_CADENCE.md resumen ejecutivo",
    ),
    (
        4, "04-data-model-er", "Data Model (PostGIS)",
        "er",
        "dji_parcels (aggregate root) — todo converge ahí",
        "5 tablas núcleo de PostGIS con sus campos clave, PKs/FKs y cardinalidades. `dji_parcels` es el aggregate root con relación 1:N a `dji_flights` y `dji_fumigations`, y 1:1 con `dji_fumigation_schedule`. `djiag_health` es singleton (1 sola fila) sin FKs.",
        "docs/AEROADMIN-AFM-OVERVIEW.md §2",
    ),
    (
        5, "05-auth-flow-sequence", "Auth Flow (NextAuth v5 + RBAC)",
        "sequence",
        "Set-Cookie de la respuesta exitosa (headline success, accent green)",
        "Secuencia del login con NextAuth v5: User → Next.js Server → NextAuth → Postgres. `alt` fragment con la rama de error (credenciales inválidas, en concern red) y la rama de éxito (Set-Cookie + 302 /). La segunda parte muestra cómo `auth()` valida el JWT en cada request posterior y aplica el role gate.",
        "docs/AEROADMIN-AFM-OVERVIEW.md §6",
    ),
    (
        6, "06-rbac-matrix", "RBAC Matrix",
        "custom grid (RBAC matrix)",
        "admin (focal, accent border) — único rol con acceso a /devices y /admin/*",
        "Grilla 9 páginas × 3 roles con la acción semántica del server-side gate por celda: `view` (allow, verde), `redirect(\"/\")` (ámbar), `notFound() 404` (rojo), `→ /login` (gris, sin sesión). Reemplaza la tabla markdown de `docs/AEROADMIN-AFM-OVERVIEW.md §6`.",
        "Reemplaza la tabla de docs/AEROADMIN-AFM-OVERVIEW.md §6",
    ),
    (
        7, "07-page-hierarchy-tree", "Page Hierarchy (Next.js App Router)",
        "tree",
        "/parcelas/[id] — única ruta con anidamiento real",
        "Jerarquía 3 niveles de `app/`: root → 5 grupos (Home, Map, Parcels, History, Admin) + 1 hoja Auth (/login) → 10 páginas tier 2 → 1 nieto (/parcelas/[id]/timeline). /devices y /admin/orphan-fumigations marcados como role-gated (admin only). /history marcado como DEPRECATED.",
        "docs/AEROADMIN-AFM-OVERVIEW.md §3",
    ),
]


# -----------------------------------------------------------------------------
# Helpers
# -----------------------------------------------------------------------------
def shade_cell(cell, hex_color):
    """Apply a background fill to a table cell (e.g. '0B5F2D')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_horizontal_rule(paragraph, color_hex="D2DDD6"):
    """Add a thin horizontal rule (bottom border) to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_paragraph_spacing(paragraph, before=0, after=6, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# -----------------------------------------------------------------------------
# Document construction
# -----------------------------------------------------------------------------
def build():
    doc = Document()

    # Page setup: A4 portrait, narrow-ish margins to maximize diagram space
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = INK

    # ------------------------------------------------------------------ COVER
    cover_top = doc.add_paragraph()
    set_paragraph_spacing(cover_top, before=120, after=0)
    eyebrow = cover_top.add_run("AeroAdmin AFM · Documentación de diagramas")
    eyebrow.font.name = "Consolas"
    eyebrow.font.size = Pt(9)
    eyebrow.font.color.rgb = SOFT
    eyebrow.font.bold = True
    eyebrow.font.all_caps = True

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=12, after=8)
    title_run = title.add_run("Diagramas del sistema")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = BRAND_GREEN
    title_run.bold = True

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, before=0, after=24, line=1.3)
    sub_run = subtitle.add_run(
        "Siete diagramas editoriales del sistema AeroAdmin AFM, "
        "generados con la skill cathrynlavery/diagram-design v2.2 y "
        "la brand skin del proyecto (verde DJI Agras como accent, "
        "alert red como concern per-node)."
    )
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = MUTED

    # Cover metadata block (small table)
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    meta.columns[0].width = Cm(4.0)
    meta.columns[1].width = Cm(13.0)
    meta_data = [
        ("Proyecto",      "AeroAdmin AFM — Valle del Cauca, Colombia"),
        ("Stack",         "Next.js 16 + React 19 + TypeScript + MapLibre 6.0 + PostGIS 3.4"),
        ("Versión",       date.today().strftime("%Y-%m-%d") + " · sprint S6"),
        ("Mantenedor",    "@agFab (single contributor)"),
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta.rows[i]
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(13.0)
        kp = row.cells[0].paragraphs[0]
        kr = kp.add_run(k)
        kr.font.name = "Consolas"
        kr.font.size = Pt(9)
        kr.font.color.rgb = SOFT
        kr.bold = True
        kp.paragraph_format.space_after = Pt(0)
        vp = row.cells[1].paragraphs[0]
        vr = vp.add_run(v)
        vr.font.size = Pt(11)
        vr.font.color.rgb = INK
        vp.paragraph_format.space_after = Pt(0)

    add_horizontal_rule(doc.add_paragraph(), "0B5F2D")

    # ----------------------------------------------------------------- TOC
    toc_intro = doc.add_paragraph()
    set_paragraph_spacing(toc_intro, before=8, after=12)
    toc_intro.add_run("Contenido").bold = True

    toc_items = [
        ("1.", "Resumen y catálogo",                      False),
        ("2.", "01 — System Architecture",                  False),
        ("3.", "02 — DJI Data Pipeline",                    False),
        ("4.", "03 — Fumigation Cadence State Machine",     False),
        ("5.", "04 — Data Model (PostGIS)",                 False),
        ("6.", "05 — Auth Flow (NextAuth v5 + RBAC)",       False),
        ("7.", "06 — RBAC Matrix",                          False),
        ("8.", "07 — Page Hierarchy (Next.js App Router)",  False),
        ("A.", "Apéndice A — Imágenes reusables (paths)",   False),
        ("B.", "Apéndice B — Cómo extender el set",          False),
    ]
    for num, label, _ in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(3)
        n_run = p.add_run(f"{num}  ")
        n_run.font.name = "Consolas"
        n_run.font.size = Pt(10)
        n_run.font.color.rgb = SOFT
        n_run.bold = True
        l_run = p.add_run(label)
        l_run.font.size = Pt(11)
        l_run.font.color.rgb = INK

    add_page_break(doc)

    # ------------------------------------------------- 1. Resumen y catálogo
    h1 = doc.add_paragraph()
    set_paragraph_spacing(h1, before=0, after=8)
    h1_run = h1.add_run("1.  Resumen y catálogo")
    h1_run.font.size = Pt(20)
    h1_run.font.color.rgb = INK
    h1_run.bold = True
    add_horizontal_rule(h1, "0B5F2D")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=10, line=1.4)
    p.add_run(
        "Los 7 diagramas cubren el sistema a 2 niveles de detalle: "
        "3 de overview (architecture, pipeline, state machine) y "
        "4 de detalle (data model, auth flow, RBAC matrix, page hierarchy). "
        "Todos están en HTML self-contained con inline SVG — se abren en "
        "cualquier browser sin build step, y los PNGs reusables están en "
    )
    p.add_run("`docs/diagrams/assets/`").font.name = "Consolas"
    p.add_run(" (ver Apéndice A).")

    cat = doc.add_table(rows=len(DIAGRAMS) + 1, cols=4)
    cat.alignment = WD_TABLE_ALIGNMENT.LEFT
    cat.autofit = False
    widths = [Cm(1.2), Cm(5.0), Cm(3.0), Cm(8.0)]
    for col_idx, w in enumerate(widths):
        cat.columns[col_idx].width = w

    hdr = cat.rows[0]
    for i, label in enumerate(["#", "Diagrama", "Tipo", "Qué muestra"]):
        cell = hdr.cells[i]
        cell.width = widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, "0B5F2D")
        run = cell.paragraphs[0].add_run(label)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    for i, (num, slug, title, dtype, focus, desc, wired) in enumerate(DIAGRAMS, start=1):
        row = cat.rows[i]
        for col_idx, w in enumerate(widths):
            row.cells[col_idx].width = w
        # # col
        c0 = row.cells[0]
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        r0 = c0.paragraphs[0].add_run(f"{num:02d}")
        r0.font.name = "Consolas"
        r0.font.size = Pt(11)
        r0.font.color.rgb = SOFT
        r0.bold = True
        c0.paragraphs[0].paragraph_format.space_after = Pt(0)
        # title col
        c1 = row.cells[1]
        r1 = c1.paragraphs[0].add_run(title)
        r1.font.size = Pt(10)
        r1.font.color.rgb = INK
        r1.bold = True
        c1.paragraphs[0].paragraph_format.space_after = Pt(0)
        # type col
        c2 = row.cells[2]
        r2 = c2.paragraphs[0].add_run(dtype)
        r2.font.name = "Consolas"
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = MUTED
        c2.paragraphs[0].paragraph_format.space_after = Pt(0)
        # desc col
        c3 = row.cells[3]
        r3 = c3.paragraphs[0].add_run(desc)
        r3.font.size = Pt(9.5)
        r3.font.color.rgb = INK
        c3.paragraphs[0].paragraph_format.space_after = Pt(0)

    add_page_break(doc)

    # ----------------------------------------------------- 2-8. Diagramas
    for num, slug, title, dtype, focus, desc, wired in DIAGRAMS:
        # Section heading with number prefix
        h2 = doc.add_paragraph()
        set_paragraph_spacing(h2, before=0, after=4)
        prefix = h2.add_run(f"{num:02d}  ")
        prefix.font.name = "Consolas"
        prefix.font.size = Pt(20)
        prefix.font.color.rgb = BRAND_GREEN
        prefix.bold = True
        tt = h2.add_run(title)
        tt.font.size = Pt(20)
        tt.font.color.rgb = INK
        tt.bold = True
        add_horizontal_rule(h2, "0B5F2D")

        # Metadata table (type, focus, wired in)
        meta = doc.add_table(rows=3, cols=2)
        meta.autofit = False
        meta.columns[0].width = Cm(2.5)
        meta.columns[1].width = Cm(14.5)
        for i, (k, v) in enumerate([("Tipo", dtype), ("Foco", focus), ("Wireado en", wired)]):
            row = meta.rows[i]
            row.cells[0].width = Cm(2.5)
            row.cells[1].width = Cm(14.5)
            kp = row.cells[0].paragraphs[0]
            kp.paragraph_format.space_after = Pt(0)
            kr = kp.add_run(k)
            kr.font.name = "Consolas"
            kr.font.size = Pt(8.5)
            kr.font.color.rgb = SOFT
            kr.bold = True
            vp = row.cells[1].paragraphs[0]
            vp.paragraph_format.space_after = Pt(0)
            vr = vp.add_run(v)
            vr.font.name = "Consolas" if k == "Tipo" else "Calibri"
            vr.font.size = Pt(10)
            vr.font.color.rgb = INK

        # Description paragraph
        d = doc.add_paragraph()
        set_paragraph_spacing(d, before=10, after=10, line=1.4)
        d_run = d.add_run(desc)
        d_run.font.size = Pt(11)
        d_run.font.color.rgb = INK

        # Image
        img_path = ASSETS / f"{slug}.png"
        if img_path.exists():
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.paragraph_format.space_before = Pt(6)
            img_p.paragraph_format.space_after = Pt(0)
            img_run = img_p.add_run()
            # Constrain to ~16cm wide so it fits A4 with 2cm margins
            img_run.add_picture(str(img_path), width=Cm(16.5))
        else:
            err = doc.add_paragraph()
            err_run = err.add_run(f"⚠ PNG no encontrado: {img_path}")
            err_run.font.color.rgb = ALERT_RED

        # Image caption (small, centered)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(0)
        cap_run = cap.add_run(f"Figura {num} — {title} · docs/diagrams/{slug}.html")
        cap_run.font.name = "Consolas"
        cap_run.font.size = Pt(8.5)
        cap_run.font.color.rgb = SOFT
        cap_run.italic = True

        add_page_break(doc)

    # ------------------------------------------- A. Apéndice: paths reusables
    hA = doc.add_paragraph()
    set_paragraph_spacing(hA, before=0, after=4)
    prefixA = hA.add_run("A.  ")
    prefixA.font.name = "Consolas"
    prefixA.font.size = Pt(20)
    prefixA.font.color.rgb = BRAND_GREEN
    prefixA.bold = True
    tA = hA.add_run("Imágenes reusables (paths)")
    tA.font.size = Pt(20)
    tA.font.color.rgb = INK
    tA.bold = True
    add_horizontal_rule(hA, "0B5F2D")

    pA = doc.add_paragraph()
    set_paragraph_spacing(pA, before=8, after=10, line=1.4)
    pA.add_run(
        "Cada diagrama tiene un PNG de alta resolución (2560×~1900) en "
    )
    pA.add_run("`docs/diagrams/assets/`").font.name = "Consolas"
    pA.add_run(
        ". Estos PNGs son los assets reusables para otros documentos "
        "(slides, otros .docx, READMEs en otros repos, etc). "
        "Los HTMLs originales son la versión canónica — los PNGs son "
        "derivados para embebido."
    )

    # Paths table
    paths = doc.add_table(rows=len(DIAGRAMS) + 2, cols=3)
    paths.alignment = WD_TABLE_ALIGNMENT.LEFT
    paths.autofit = False
    widths_p = [Cm(1.2), Cm(8.5), Cm(7.3)]
    for col_idx, w in enumerate(widths_p):
        paths.columns[col_idx].width = w

    hdr_p = paths.rows[0]
    for i, label in enumerate(["#", "Path relativo al repo", "Path absoluto (Windows)"]):
        cell = hdr_p.cells[i]
        cell.width = widths_p[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, "0B5F2D")
        r = cell.paragraphs[0].add_run(label)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.bold = True
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    for i, (num, slug, *_rest) in enumerate(DIAGRAMS, start=1):
        row = paths.rows[i]
        for col_idx, w in enumerate(widths_p):
            row.cells[col_idx].width = w
        c0 = row.cells[0]
        r0 = c0.paragraphs[0].add_run(f"{num:02d}")
        r0.font.name = "Consolas"
        r0.font.size = Pt(11)
        r0.font.color.rgb = SOFT
        r0.bold = True
        c0.paragraphs[0].paragraph_format.space_after = Pt(0)
        c1 = row.cells[1]
        r1 = c1.paragraphs[0].add_run(f"docs/diagrams/assets/{slug}.png")
        r1.font.name = "Consolas"
        r1.font.size = Pt(9)
        r1.font.color.rgb = INK
        c1.paragraphs[0].paragraph_format.space_after = Pt(0)
        c2 = row.cells[2]
        r2 = c2.paragraphs[0].add_run(str(ASSETS / f"{slug}.png"))
        r2.font.name = "Consolas"
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = MUTED
        c2.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Final note row
    note_row = paths.rows[-1]
    note_row.cells[0].merge(note_row.cells[2])
    np_para = note_row.cells[0].paragraphs[0]
    np_para.paragraph_format.space_after = Pt(0)
    np_run = np_para.add_run(
        "Tip: insertar en otro docx → "
    )
    np_run.font.size = Pt(9.5)
    np_run.font.color.rgb = MUTED
    code_run = np_para.add_run("Insert → Picture → From File")
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(9.5)
    code_run.font.color.rgb = INK
    np_para.add_run(" y seleccionar el PNG.").font.size = Pt(9.5)

    # ------------------------------------------- B. Apéndice: cómo extender
    hB = doc.add_paragraph()
    set_paragraph_spacing(hB, before=18, after=4)
    prefixB = hB.add_run("B.  ")
    prefixB.font.name = "Consolas"
    prefixB.font.size = Pt(20)
    prefixB.font.color.rgb = BRAND_GREEN
    prefixB.bold = True
    tB = hB.add_run("Cómo extender el set")
    tB.font.size = Pt(20)
    tB.font.color.rgb = INK
    tB.bold = True
    add_horizontal_rule(hB, "0B5F2D")

    refs = [
        ("README.md", "Catálogo y procedimiento paso a paso para agregar un nuevo diagrama."),
        ("HANDOFF.md", "Contexto, decisiones, errores que ya pagué — para que otro agente no re-descubra."),
        ("style-guide.md", "Brand skin project-owned (paleta, tipografía, treatment de nodos)."),
    ]
    for fname, desc in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        code = p.add_run(f"docs/diagrams/{fname}")
        code.font.name = "Consolas"
        code.font.size = Pt(10)
        code.font.color.rgb = BRAND_GREEN
        code.bold = True
        p.add_run(" — ").font.color.rgb = MUTED
        d = p.add_run(desc)
        d.font.size = Pt(10.5)
        d.font.color.rgb = INK

    tip = doc.add_paragraph()
    set_paragraph_spacing(tip, before=10, after=4, line=1.4)
    tip_run = tip.add_run("Render utility: ")
    tip_run.font.size = Pt(10.5)
    tip_run.font.color.rgb = INK
    tip_run.bold = True
    tip_code = tip.add_run("node scripts/render-diagram.js <input.html> <output.png>")
    tip_code.font.name = "Consolas"
    tip_code.font.size = Pt(10)
    tip_code.font.color.rgb = BRAND_GREEN
    tip.add_run(
        "  ·  verificar visualmente cada PNG antes de commitear."
    ).font.size = Pt(10.5)

    # ----------------------------------------------------------------- Save
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK -> {OUT}")
    print(f"     size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    build()
